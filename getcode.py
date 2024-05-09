import os
from pathlib import Path

import docx
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.shared import Pt, Cm

EXCLUDE_LIST = ["cvmatandqimage.cpp", "cvmatandqimage.h", "imageutil.h", "popen2.h"]

lab_path = "."

doc = docx.Document()

first = True

for root, dirs, files in os.walk(lab_path):
    if str(Path(root).relative_to(lab_path)).startswith("build"):
        continue
    for file in files:
        if (file.endswith(".cpp") or file.endswith(".h")) and file not in EXCLUDE_LIST:
            with open(os.path.join(root, file), "r", encoding="utf-8") as f:
                head, tail = os.path.split(f.name)
                if not first:
                    tmp_para = doc.add_paragraph()
                    tmp_para.paragraph_format.space_after = Pt(0)
                    tmp_para.paragraph_format.first_line_indent = Cm(1.25)
                title_para = doc.add_paragraph()
                title_run = title_para.add_run("Файл " + tail + ":")
                title_run.font.size = Pt(14)
                title_run.font.name = "Times New Roman"
                title_para.paragraph_format.line_spacing = 1
                title_para.paragraph_format.space_after = Pt(0)
                title_para.paragraph_format.first_line_indent = Cm(1.25)
                for line in f.read().rstrip().split("\n"):
                    para = doc.add_paragraph()
                    para_run = para.add_run(line)
                    para_run.font.size = Pt(10)
                    para_run.font.name = "Courier New"
                    para.paragraph_format.line_spacing = 1
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.first_line_indent = Cm(1.25)
                first = False
doc.save("code.docx")
